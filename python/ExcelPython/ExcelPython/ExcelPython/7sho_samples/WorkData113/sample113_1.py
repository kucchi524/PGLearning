import openpyxl
import docx
fname = '販売記録08.xlsx'
wb = openpyxl.load_workbook(fname)
ws = wb.active
doc = docx.Document('販売報告01.docx')
month = doc.paragraphs[3].text
ws['D2'].value = month[1:3]
person = doc.paragraphs[2].text
ws['E2'].value = person[3:]
wb.save(fname)
