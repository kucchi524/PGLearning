import openpyxl
import docx
fname = '販売報告02.docx'
wb = openpyxl.load_workbook('販売記録09.xlsx')
ws = wb.active
doc = docx.Document(fname)
tinfo = doc.paragraphs[3].text
doc.paragraphs[3].text = tinfo[:1] +ws['C2'].value + '：' + tinfo[1:]
tbl = doc.tables[0]
for row in ws.iter_rows(min_row=5):
    nrow = tbl.add_row()
    for wcell, xcell in zip(nrow.cells, row):
        wcell.text = str(xcell.value)
doc.save(fname)
